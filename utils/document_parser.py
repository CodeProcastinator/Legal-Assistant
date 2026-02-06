"""
Document Parser Module
Handles parsing of PDF, DOCX, and TXT files for contract analysis
"""

import re
import io
from typing import Dict, List, Optional, Tuple
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False


class DocumentParser:
    """
    A comprehensive document parser for legal contracts.
    Supports PDF, DOCX, and TXT formats.
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        self.last_error = None
        self.metadata = {}
    
    def parse(self, file_content: bytes, filename: str) -> Dict:
        """
        Parse a document and extract text content.
        
        Args:
            file_content: Raw bytes of the file
            filename: Name of the file with extension
            
        Returns:
            Dict containing extracted text, metadata, and status
        """
        self.last_error = None
        self.metadata = {}
        
        extension = Path(filename).suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            return {
                'success': False,
                'error': f'Unsupported file format: {extension}. Supported formats: PDF, DOCX, TXT',
                'text': '',
                'metadata': {}
            }
        
        try:
            if extension == '.pdf':
                text = self._parse_pdf(file_content)
            elif extension in ['.docx', '.doc']:
                text = self._parse_docx(file_content)
            elif extension == '.txt':
                text = self._parse_txt(file_content)
            else:
                text = ''
            
            # Clean and normalize text
            cleaned_text = self._clean_text(text)
            
            # Extract metadata
            self.metadata['filename'] = filename
            self.metadata['extension'] = extension
            self.metadata['char_count'] = len(cleaned_text)
            self.metadata['word_count'] = len(cleaned_text.split())
            self.metadata['page_estimate'] = max(1, len(cleaned_text) // 3000)
            
            return {
                'success': True,
                'text': cleaned_text,
                'metadata': self.metadata,
                'error': None
            }
            
        except Exception as e:
            self.last_error = str(e)
            return {
                'success': False,
                'error': f'Error parsing document: {str(e)}',
                'text': '',
                'metadata': {}
            }
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF file content."""
        text_parts = []
        
        # Try pdfplumber first (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    self.metadata['page_count'] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                if text_parts:
                    return '\n\n'.join(text_parts)
            except Exception:
                pass
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                self.metadata['page_count'] = len(pdf_reader.pages)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return '\n\n'.join(text_parts)
            except Exception as e:
                raise Exception(f"PDF parsing failed: {str(e)}")
        
        raise Exception("No PDF parsing library available. Install PyPDF2 or pdfplumber.")
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX file content."""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            self.metadata['paragraph_count'] = len(doc.paragraphs)
            self.metadata['table_count'] = len(doc.tables)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {str(e)}")
    
    def _parse_txt(self, content: bytes) -> str:
        """Parse plain text file content."""
        # Detect encoding
        if CHARDET_AVAILABLE:
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        else:
            encoding = 'utf-8'
        
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            # Fallback encodings
            for enc in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
                try:
                    return content.decode(enc)
                except UnicodeDecodeError:
                    continue
            raise Exception("Could not decode text file with any known encoding")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = re.sub(r'\t+', ' ', text)
        
        # Remove page numbers (common patterns)
        text = re.sub(r'\n\s*Page\s+\d+\s*(?:of\s+\d+)?\s*\n', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\s*-\s*\d+\s*-\s*\n', '\n', text)
        
        # Fix common OCR issues
        text = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', text)  # Missing spaces
        
        # Normalize quotes and dashes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        text = text.replace('–', '-').replace('—', '-')
        
        return text.strip()
    
    def extract_sections(self, text: str) -> List[Dict]:
        """
        Extract document sections based on common contract patterns.
        
        Args:
            text: Cleaned document text
            
        Returns:
            List of section dictionaries with title and content
        """
        sections = []
        
        # Common section header patterns
        patterns = [
            r'(?:^|\n)(\d+\.?\s+[A-Z][A-Za-z\s]+)(?:\n|:)',  # 1. DEFINITIONS
            r'(?:^|\n)(ARTICLE\s+[IVXLCDM\d]+[:\s]+[A-Za-z\s]+)',  # ARTICLE I: Scope
            r'(?:^|\n)(CLAUSE\s+\d+[:\s]+[A-Za-z\s]+)',  # CLAUSE 1: Terms
            r'(?:^|\n)([A-Z][A-Z\s]+)(?:\n)',  # ALL CAPS HEADERS
            r'(?:^|\n)(\d+\.\d+\s+[A-Za-z][A-Za-z\s]+)',  # 1.1 Sub-sections
        ]
        
        all_matches = []
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                all_matches.append({
                    'title': match.group(1).strip(),
                    'start': match.start(),
                    'end': match.end()
                })
        
        # Sort by position
        all_matches.sort(key=lambda x: x['start'])
        
        # Extract content between sections
        for i, match in enumerate(all_matches):
            start = match['end']
            end = all_matches[i + 1]['start'] if i + 1 < len(all_matches) else len(text)
            content = text[start:end].strip()
            
            if content and len(content) > 20:  # Filter out empty sections
                sections.append({
                    'title': match['title'],
                    'content': content[:5000],  # Limit content length
                    'position': i + 1
                })
        
        return sections
    
    def get_document_type_hints(self, text: str) -> Dict:
        """
        Analyze text to provide hints about document type.
        
        Args:
            text: Document text
            
        Returns:
            Dict with document type hints and confidence scores
        """
        text_lower = text.lower()
        
        type_indicators = {
            'employment': [
                'employment agreement', 'employee', 'employer', 'salary',
                'termination of employment', 'probation period', 'working hours',
                'leave policy', 'notice period', 'resignation'
            ],
            'vendor': [
                'vendor agreement', 'supplier', 'purchase order', 'delivery',
                'payment terms', 'invoice', 'goods', 'supply chain'
            ],
            'service': [
                'service agreement', 'service provider', 'scope of services',
                'service level', 'sla', 'deliverables', 'milestone'
            ],
            'lease': [
                'lease agreement', 'landlord', 'tenant', 'rent', 'premises',
                'security deposit', 'lease term', 'rental'
            ],
            'nda': [
                'non-disclosure', 'confidential information', 'confidentiality',
                'trade secret', 'proprietary information', 'nda'
            ],
            'partnership': [
                'partnership agreement', 'partner', 'profit sharing',
                'capital contribution', 'joint venture', 'partnership deed'
            ],
            'loan': [
                'loan agreement', 'borrower', 'lender', 'principal amount',
                'interest rate', 'repayment', 'emi', 'collateral'
            ],
            'consulting': [
                'consulting agreement', 'consultant', 'advisory', 'retainer',
                'consulting services', 'independent contractor'
            ]
        }
        
        scores = {}
        for doc_type, keywords in type_indicators.items():
            score = sum(1 for kw in keywords if kw in text_lower)
            scores[doc_type] = min(score / len(keywords), 1.0)
        
        # Get best match
        best_type = max(scores, key=scores.get) if scores else 'general'
        best_score = scores.get(best_type, 0)
        
        return {
            'likely_type': best_type if best_score > 0.1 else 'general',
            'confidence': best_score,
            'all_scores': scores
        }


# For testing
if __name__ == "__main__":
    parser = DocumentParser()
    
    # Test with sample text
    sample = b"""
    EMPLOYMENT AGREEMENT
    
    This Employment Agreement is entered into between ABC Company (Employer) 
    and John Doe (Employee) on January 1, 2024.
    
    1. POSITION AND DUTIES
    The Employee shall serve as Software Engineer with responsibilities as defined.
    
    2. COMPENSATION
    The Employee shall receive a monthly salary of Rs. 50,000.
    
    3. TERMINATION
    Either party may terminate this agreement with 30 days written notice.
    """
    
    result = parser.parse(sample, "test_contract.txt")
    print(f"Success: {result['success']}")
    print(f"Word Count: {result['metadata'].get('word_count', 0)}")
    
    type_hints = parser.get_document_type_hints(result['text'])
    print(f"Document Type: {type_hints['likely_type']} (confidence: {type_hints['confidence']:.2f})")
